import docx
import sys
import io

def docx_to_markdown(docx_path, md_path):
    doc = docx.Document(docx_path)
    md_lines = []
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        
        # Check for headings based on style or numberings
        if p.style.name.startswith('Heading 1'):
            md_lines.append(f"# {text}\n")
        elif p.style.name.startswith('Heading 2'):
            md_lines.append(f"## {text}\n")
        elif p.style.name.startswith('Heading 3'):
            md_lines.append(f"### {text}\n")
        else:
            md_lines.append(f"{text}\n")
            
    # Also extract tables
    for i, table in enumerate(doc.tables):
        md_lines.append(f"\n### Bảng {i+1}\n")
        for row_idx, row in enumerate(table.rows):
            row_texts = []
            for cell in row.cells:
                # Get cell text, strip duplicates if openxml creates nested cells
                cell_text = " ".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
                row_texts.append(cell_text.replace("\n", "<br>"))
            
            if row_idx == 0:
                md_lines.append("| " + " | ".join(row_texts) + " |")
                md_lines.append("| " + " | ".join(["---"] * len(row_texts)) + " |")
            else:
                md_lines.append("| " + " | ".join(row_texts) + " |")
                
    with io.open(md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_lines))

if __name__ == "__main__":
    docx_to_markdown(
        'DAN_CS445_AE/0225.CS445-AE.8.1.ProjectTestPlanSprint1.docx', 
        'Sample_Test_Plan_Reference.md'
    )
    print("Done converting DOCX to Markdown!")
